from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"D:\website kasir")
OUT = ROOT / "output" / "manual"
IMG = OUT / "screenshots"
DOCX = OUT / "Technical_User_Manual_Admin_Solutions_Inovatif_POS.docx"

MONTHS_ID = {
    1: "Januari",
    2: "Februari",
    3: "Maret",
    4: "April",
    5: "Mei",
    6: "Juni",
    7: "Juli",
    8: "Agustus",
    9: "September",
    10: "Oktober",
    11: "November",
    12: "Desember",
}


def tanggal_indonesia(dt):
    return f"{dt.day} {MONTHS_ID[dt.month]} {dt.year}"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="111827"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.size = Pt(9)


def style_table(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
            if row_i == 0 and header:
                set_cell_shading(cell, "E8EEF5")


def add_caption(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"Gambar {number}. {title}")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(71, 85, 105)


def add_figure(doc, number, filename, caption, width=6.45):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(IMG / filename), width=Inches(width))
    add_caption(doc, number, caption)


def add_note(doc, title, text, tone="info"):
    colors = {
        "info": ("E0F2FE", "075985"),
        "warning": ("FEF3C7", "92400E"),
        "success": ("DCFCE7", "166534"),
    }
    fill, color = colors.get(tone, colors["info"])
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + " ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.size = Pt(9.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor(51, 65, 85)
    doc.add_paragraph()


def add_steps(doc, steps):
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(step)


def add_bullets(doc, bullets):
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_module(doc, title, purpose, access, steps, notes, fig_no, fig_file, fig_caption):
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    p.add_run("Tujuan modul. ").bold = True
    p.add_run(purpose)
    p = doc.add_paragraph()
    p.add_run("Akses pengguna. ").bold = True
    p.add_run(access)
    add_figure(doc, fig_no, fig_file, fig_caption)
    doc.add_heading("Langkah Penggunaan", level=2)
    add_steps(doc, steps)
    if notes:
        doc.add_heading("Catatan Operasional", level=2)
        add_bullets(doc, notes)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, "0F766E", 14, 7),
        ("Heading 2", 13, "0F766E", 10, 5),
        ("Heading 3", 11.5, "1F2937", 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Number", "List Bullet"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("Technical User Manual")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(15, 118, 110)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Aplikasi Admin Solutions Inovatif - Sistem Manajemen Bisnis & POS")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(15, 23, 42)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Panduan pengguna teknis untuk operasional toko, kasir, inventaris, laporan, akuntansi, dan administrasi platform.\n")
    meta.add_run(f"Versi dokumen: 1.0 | Tanggal: {tanggal_indonesia(datetime.now())}")

    logo = ROOT / "public" / "logo.jpg"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Inches(1.1))

    add_note(
        doc,
        "Ruang lingkup.",
        "Dokumen ini disusun berdasarkan tampilan dan perilaku aplikasi yang berjalan di lingkungan lokal pada saat pemeriksaan. Setiap modul dijelaskan sesuai menu, form, tombol, dan alur yang tersedia di aplikasi.",
        "info",
    )
    doc.add_page_break()

    # TOC
    doc.add_heading("Daftar Isi", level=1)
    for item in [
        "1. Pendahuluan",
        "2. Peran Pengguna dan Hak Akses",
        "3. Masuk ke Aplikasi",
        "4. Navigasi Umum",
        "5. Dashboard",
        "6. Kasir (POS)",
        "7. Produk",
        "8. Kategori",
        "9. Inventaris",
        "10. Transaksi",
        "11. Pelanggan",
        "12. Pemasok",
        "13. Laporan",
        "14. Akuntansi",
        "15. Karyawan",
        "16. Pengaturan",
        "17. Platform Admin",
        "18. Praktik Penggunaan yang Disarankan",
    ]:
        doc.add_paragraph(item)
    doc.add_page_break()

    doc.add_heading("Pendahuluan", level=1)
    doc.add_paragraph(
        "Admin Solutions Inovatif adalah aplikasi POS dan manajemen bisnis untuk membantu UMKM mengelola penjualan, produk, stok, pelanggan, pemasok, laporan, biaya operasional, dan pengaturan toko dalam satu sistem. Aplikasi menggunakan konsep tenant, sehingga data toko dipisahkan berdasarkan akun bisnis yang sedang digunakan."
    )
    doc.add_paragraph(
        "Manual ini ditujukan untuk pengguna operasional, pemilik toko, dan administrator platform. Fokus panduan adalah cara menggunakan fitur dari antarmuka aplikasi, bukan konfigurasi teknis server atau pengembangan kode."
    )

    doc.add_heading("Ringkasan Modul", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Modul", "Fungsi Utama", "Pengguna Umum"]):
        set_cell_text(cell, text, True)
    rows = [
        ("Dashboard", "Memantau KPI penjualan, transaksi terbaru, dan peringatan stok.", "Owner, Cashier"),
        ("Kasir (POS)", "Memilih produk, mengelola keranjang, menerima pembayaran, dan mencetak struk.", "Owner, Cashier"),
        ("Produk & Kategori", "Mengelola katalog, harga, stok awal, gambar produk, dan pengelompokan kategori.", "Owner"),
        ("Inventaris", "Mencatat stok masuk, stok keluar, dan penyesuaian stok.", "Owner"),
        ("Transaksi", "Melihat riwayat transaksi, detail struk, dan melakukan refund bila berwenang.", "Owner, Cashier"),
        ("Laporan & Akuntansi", "Menganalisis penjualan, laba rugi, neraca, dan biaya operasional.", "Owner"),
        ("Karyawan & Pengaturan", "Mengelola pengguna toko, profil toko, pajak, logo, dan metode pembayaran.", "Owner"),
        ("Platform Admin", "Memantau tenant, paket layanan, status toko, dan statistik platform.", "Super Admin"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    style_table(table)

    doc.add_heading("Peran Pengguna dan Hak Akses", level=1)
    doc.add_paragraph("Aplikasi membedakan fitur berdasarkan peran pengguna. Menu yang tampil di sidebar akan menyesuaikan hak akses akun yang sedang login.")
    table = doc.add_table(rows=1, cols=4)
    for cell, text in zip(table.rows[0].cells, ["Peran", "Fokus Tugas", "Akses Utama", "Pembatasan"]):
        set_cell_text(cell, text, True)
    role_rows = [
        ("Super Admin", "Mengelola platform", "Platform Admin, statistik tenant, paket layanan, status tenant", "Tidak digunakan untuk operasional kasir harian toko tertentu."),
        ("Owner", "Mengelola toko", "Semua modul operasional toko, laporan, akuntansi, karyawan, dan pengaturan", "Hanya melihat data tenant/toko miliknya."),
        ("Cashier", "Memproses penjualan", "Dashboard, POS, produk, kategori, inventaris, transaksi, pelanggan", "Tidak mengakses laporan, akuntansi, karyawan, pemasok, dan pengaturan."),
    ]
    for row in role_rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    style_table(table)

    doc.add_heading("Masuk ke Aplikasi", level=1)
    doc.add_paragraph("Pengguna memulai aktivitas dari halaman login. Masukkan alamat email dan kata sandi yang diberikan oleh administrator, lalu tekan tombol Masuk. Pengguna baru dapat memilih Daftar sekarang untuk membuat akun owner dan tenant baru apabila fitur registrasi diaktifkan.")
    add_figure(doc, 1, "01-login.png", "Halaman login aplikasi")
    doc.add_heading("Langkah Login", level=2)
    add_steps(doc, [
        "Buka alamat aplikasi dari browser.",
        "Isi Alamat Email sesuai akun pengguna.",
        "Isi Kata Sandi. Gunakan ikon mata untuk menampilkan atau menyembunyikan kata sandi bila diperlukan.",
        "Klik Masuk. Jika kredensial valid, aplikasi mengarahkan pengguna ke Dashboard.",
        "Jika muncul pesan kesalahan, periksa kembali email, kata sandi, atau status akun pengguna.",
    ])
    add_note(doc, "Peringatan.", "Jangan membagikan akun atau kata sandi antar pengguna. Setiap transaksi menyimpan informasi kasir sehingga akun harus digunakan oleh pemiliknya sendiri.", "warning")

    doc.add_heading("Navigasi Umum", level=1)
    doc.add_paragraph("Setelah login, aplikasi menampilkan sidebar di sisi kiri, informasi pengguna, indikator tenant terisolasi untuk akun toko, top bar, dan ikon notifikasi. Klik menu sidebar untuk berpindah modul. Tombol Keluar berada di bagian bawah sidebar.")
    add_figure(doc, 2, "02-dashboard.png", "Dashboard utama dan navigasi aplikasi")
    add_bullets(doc, [
        "Sidebar menampilkan modul sesuai hak akses pengguna.",
        "Top bar menampilkan nama halaman aktif dan tombol notifikasi.",
        "Panel notifikasi digunakan untuk informasi stok menipis, transaksi, pembayaran, dan refund.",
        "Pada layar kecil, sidebar dibuka melalui tombol menu di kiri atas.",
    ])

    add_module(
        doc,
        "Dashboard",
        "Dashboard memberi gambaran cepat atas performa toko, meliputi penjualan hari ini, penjualan bulan ini, jumlah produk, stok menipis, transaksi terbaru, total pelanggan, total karyawan, dan rata-rata transaksi.",
        "Owner dan Cashier dapat melihat dashboard. Data yang tampil mengikuti tenant/toko dari akun pengguna.",
        [
            "Buka menu Dashboard dari sidebar.",
            "Periksa kartu KPI untuk melihat kondisi penjualan dan stok secara cepat.",
            "Gunakan daftar Transaksi Terbaru untuk memantau aktivitas kasir terakhir.",
            "Perhatikan bagian Peringatan Stok Menipis. Produk pada bagian ini perlu diprioritaskan untuk restock atau penyesuaian stok.",
        ],
        [
            "Dashboard tidak digunakan untuk input data. Perubahan data dilakukan melalui modul POS, Produk, Inventaris, atau modul terkait.",
            "Nilai KPI akan berubah setelah transaksi, refund, atau mutasi stok berhasil tersimpan.",
        ],
        3,
        "02-dashboard.png",
        "Ringkasan KPI, transaksi terbaru, dan peringatan stok",
    )

    add_module(
        doc,
        "Kasir (POS)",
        "Modul Kasir (POS) digunakan untuk memproses penjualan. Pengguna dapat mencari produk, memasukkan produk ke keranjang, memilih metode pembayaran, menggunakan split payment, menghitung kembalian, dan mencetak struk.",
        "Owner dan Cashier dapat mengakses POS. Metode pembayaran yang muncul mengikuti konfigurasi di Pengaturan.",
        [
            "Buka menu Kasir (POS).",
            "Cari produk melalui kolom pencarian berdasarkan nama produk, SKU, atau barcode.",
            "Klik kartu produk untuk memasukkannya ke Keranjang Belanja.",
            "Gunakan tombol plus atau minus untuk mengubah jumlah item. Gunakan ikon hapus untuk mengeluarkan item dari keranjang.",
            "Periksa subtotal, pajak, dan total.",
            "Pilih metode pembayaran, misalnya Tunai, QRIS, Bank Transfer, E-Wallet, atau Split Payment bila tersedia.",
            "Isi nominal pembayaran. Untuk pembayaran tunai, sistem menampilkan kembalian.",
            "Klik Bayar & Cetak Struk setelah nominal pembayaran cukup.",
        ],
        [
            "Produk dengan stok 0 tidak dapat ditambahkan ke keranjang.",
            "Tombol bayar tidak aktif jika keranjang kosong, metode pembayaran tidak tersedia, atau nominal pembayaran kurang dari total.",
            "Setelah checkout berhasil, stok produk berkurang dan struk dapat dicetak ulang dari dialog sukses.",
        ],
        4,
        "03-pos.png",
        "Modul POS dengan produk, keranjang, pembayaran, dan tombol cetak struk",
    )

    doc.add_heading("Produk", level=1)
    doc.add_paragraph("Modul Produk berfungsi sebagai master katalog barang yang dijual. Halaman ini menampilkan pencarian, filter kategori, tabel produk, harga beli, harga jual, stok, status, dan tombol aksi untuk pengguna yang memiliki hak kelola.")
    add_figure(doc, 5, "04-products.png", "Daftar produk, pencarian, filter, dan aksi produk")
    add_steps(doc, [
        "Buka menu Produk.",
        "Gunakan kolom Cari produk untuk menemukan produk berdasarkan nama atau SKU.",
        "Gunakan filter Semua kategori untuk membatasi tampilan berdasarkan kategori.",
        "Klik Tambah Produk untuk membuat produk baru.",
        "Gunakan ikon pensil untuk mengubah produk dan ikon hapus untuk menghapus produk bila diperlukan.",
    ])
    doc.add_heading("Menambah atau Mengubah Produk", level=2)
    doc.add_paragraph("Form produk memuat nama, SKU, barcode, kategori, harga beli, harga jual, stok awal, stok minimum, gambar produk, deskripsi, dan status untuk mode edit.")
    add_figure(doc, 6, "05-product-form.png", "Form tambah atau ubah produk")
    add_steps(doc, [
        "Isi Nama Produk dan SKU sebagai data wajib.",
        "Isi Barcode bila produk memiliki kode barcode.",
        "Pilih Kategori bila produk perlu dikelompokkan.",
        "Masukkan Harga Beli dan Harga Jual.",
        "Isi Stok Awal dan Stok Minimum agar dashboard dapat memberi peringatan stok menipis.",
        "Unggah gambar produk bila tersedia.",
        "Klik Tambah Produk atau Simpan Perubahan.",
    ])

    add_module(
        doc,
        "Kategori",
        "Kategori membantu mengelompokkan produk agar pencarian, filter, dan pelaporan katalog lebih rapi.",
        "Owner dapat menambah, mengubah, dan menghapus kategori. Cashier dapat melihat kategori sesuai menu yang tampil.",
        [
            "Buka menu Kategori.",
            "Klik Tambah Kategori.",
            "Isi Nama Kategori, Keterangan, dan pilih warna kategori.",
            "Klik Tambah Kategori atau Simpan Kategori.",
            "Gunakan tombol Ubah untuk memperbarui kategori dan ikon hapus untuk menghapus kategori.",
        ],
        [
            "Jumlah produk pada kartu kategori menunjukkan berapa produk yang terhubung dengan kategori tersebut.",
            "Sebelum menghapus kategori, pastikan dampaknya terhadap produk yang sedang memakai kategori tersebut sudah dipahami.",
        ],
        7,
        "06-categories.png",
        "Daftar kategori produk",
    )

    doc.add_heading("Inventaris", level=1)
    doc.add_paragraph("Inventaris digunakan untuk mencatat riwayat perubahan stok. Sistem membedakan mutasi menjadi Stok Masuk, Stok Keluar, dan Penyesuaian. Halaman juga menampilkan indikator produk yang berada di bawah stok minimum.")
    add_figure(doc, 8, "07-inventory.png", "Riwayat mutasi stok dan indikator stok menipis")
    add_steps(doc, [
        "Buka menu Inventaris.",
        "Periksa pesan stok menipis di bagian atas halaman.",
        "Baca tabel riwayat mutasi untuk mengetahui jenis perubahan, produk, jumlah, stok saat ini, keterangan, dan tanggal.",
        "Klik Mutasi Stok untuk menambah catatan stok baru.",
    ])
    doc.add_heading("Mencatat Mutasi Stok", level=2)
    add_figure(doc, 9, "08-inventory-form.png", "Form mutasi stok")
    add_steps(doc, [
        "Pilih Jenis Mutasi: Stok Masuk, Stok Keluar, atau Penyesuaian Stok.",
        "Pilih Produk yang akan diubah stoknya.",
        "Masukkan Jumlah.",
        "Isi Catatan atau Alasan agar riwayat stok mudah diaudit.",
        "Klik Simpan Mutasi.",
    ])

    add_module(
        doc,
        "Transaksi",
        "Modul Transaksi menampilkan daftar struk penjualan, filter tanggal, status pembayaran, metode pembayaran, total transaksi, detail item, dan fungsi refund untuk pengguna berwenang.",
        "Owner dan Cashier dapat melihat riwayat transaksi. Refund tersedia untuk pengguna yang memiliki hak kelola.",
        [
            "Buka menu Transaksi.",
            "Gunakan filter Dari dan Hingga untuk membatasi periode transaksi.",
            "Klik ikon mata pada baris transaksi untuk membuka detail struk.",
            "Pada detail transaksi, periksa item, subtotal, pajak, diskon, total, metode pembayaran, kasir, pelanggan, dan status.",
            "Jika perlu dan memiliki hak akses, klik Refund Transaksi pada transaksi berstatus LUNAS.",
        ],
        [
            "Refund mengubah status transaksi dan memengaruhi ringkasan laporan.",
            "Gunakan refund hanya untuk pembatalan transaksi yang benar-benar sah secara operasional.",
        ],
        10,
        "09-transactions.png",
        "Daftar transaksi dan filter periode",
    )

    add_module(
        doc,
        "Pelanggan",
        "Modul Pelanggan menyimpan data pelanggan berupa nama, nomor telepon, email, alamat, dan jumlah transaksi yang pernah dilakukan.",
        "Owner dapat menambah, mengubah, dan menghapus pelanggan. Cashier dapat melihat data pelanggan jika menu tersedia.",
        [
            "Buka menu Pelanggan.",
            "Gunakan kolom Cari pelanggan untuk menemukan data pelanggan.",
            "Klik Tambah Pelanggan untuk membuat data baru.",
            "Isi Nama Pelanggan sebagai data wajib, lalu lengkapi nomor telepon, email, dan alamat jika tersedia.",
            "Gunakan ikon pensil untuk mengubah data dan ikon hapus untuk menghapus data.",
        ],
        [
            "Data pelanggan membantu riwayat transaksi lebih mudah ditelusuri.",
            "Pastikan data kontak diisi dengan format yang konsisten agar mudah dicari.",
        ],
        11,
        "10-customers.png",
        "Daftar pelanggan dan pencarian pelanggan",
    )

    add_module(
        doc,
        "Pemasok",
        "Modul Pemasok menyimpan data pihak pemasok barang, termasuk contact person, nomor telepon, email, alamat, dan jumlah purchase order terkait.",
        "Owner dapat mengelola data pemasok. Menu ini tidak ditampilkan untuk Cashier.",
        [
            "Buka menu Pemasok.",
            "Klik Tambah Pemasok.",
            "Isi Nama Pemasok sebagai data wajib.",
            "Lengkapi Contact Person, nomor telepon, email, dan alamat.",
            "Klik Tambah Pemasok atau Simpan Perubahan.",
            "Gunakan ikon pensil untuk mengubah data dan ikon hapus untuk menghapus pemasok.",
        ],
        [
            "Data pemasok mendukung pengelolaan pengadaan dan histori purchase order.",
            "Gunakan nama pemasok yang konsisten agar mudah dibedakan pada laporan atau proses pembelian.",
        ],
        12,
        "11-suppliers.png",
        "Daftar pemasok",
    )

    add_module(
        doc,
        "Laporan",
        "Modul Laporan menampilkan analisis penjualan berdasarkan periode, termasuk total pendapatan, laba bersih, jumlah transaksi, rata-rata transaksi, pajak terkumpul, grafik penjualan harian, produk terlaris, dan rincian metode pembayaran.",
        "Owner dapat mengakses laporan. Cashier tidak melihat menu ini.",
        [
            "Buka menu Laporan.",
            "Pilih tanggal awal dan tanggal akhir periode laporan.",
            "Periksa kartu ringkasan pendapatan, laba, transaksi, pajak, dan diskon.",
            "Gunakan grafik Penjualan Harian untuk melihat pola penjualan per hari.",
            "Lihat Produk Terlaris untuk mengetahui produk dengan performa terbaik.",
            "Lihat Metode Pembayaran untuk memantau distribusi pembayaran.",
        ],
        [
            "Periode default mengikuti bulan berjalan.",
            "Laporan bergantung pada transaksi yang berhasil dan data refund yang tercatat.",
        ],
        13,
        "12-reports.png",
        "Laporan penjualan dan analisis performa",
    )

    doc.add_heading("Akuntansi", level=1)
    doc.add_paragraph("Modul Akuntansi terdiri dari tiga tab: Laba Rugi, Neraca Keuangan, dan Biaya Operasional. Modul ini membantu owner melihat kesehatan keuangan toko dari sisi pendapatan, HPP, beban, laba, aset, liabilitas, ekuitas, dan pengeluaran.")
    add_figure(doc, 14, "13-accounting-profit-loss.png", "Tab Laba Rugi pada modul Akuntansi")
    doc.add_heading("Tab Laba Rugi", level=2)
    add_steps(doc, [
        "Pilih tab Laba Rugi.",
        "Tentukan tanggal awal dan akhir periode.",
        "Periksa Pendapatan Usaha, HPP, Laba Kotor, Beban Operasional, dan Laba Bersih.",
        "Gunakan grafik pendapatan vs beban untuk memahami proporsi biaya terhadap penjualan.",
        "Periksa rincian beban operasional untuk mengetahui kategori biaya terbesar.",
    ])
    doc.add_heading("Tab Neraca Keuangan", level=2)
    add_steps(doc, [
        "Pilih tab Neraca Keuangan.",
        "Tentukan tanggal neraca pada field Per tanggal.",
        "Periksa status Seimbang atau Belum Seimbang.",
        "Bandingkan total aset dengan total liabilitas dan ekuitas.",
    ])
    doc.add_heading("Tab Biaya Operasional", level=2)
    doc.add_paragraph("Tab Biaya Operasional digunakan untuk mencatat pengeluaran seperti sewa tempat, utilitas, gaji, pemasaran, perlengkapan, dan biaya lainnya.")
    add_figure(doc, 15, "14-accounting-expenses.png", "Tab Biaya Operasional dan tombol catat pengeluaran")
    add_steps(doc, [
        "Pilih tab Biaya Operasional.",
        "Gunakan filter tanggal untuk menentukan periode biaya.",
        "Klik Catat Pengeluaran.",
        "Pilih kategori, isi keterangan, nominal, dan tanggal.",
        "Klik Simpan Pengeluaran.",
        "Gunakan ikon hapus untuk menghapus pengeluaran yang salah input.",
    ])

    add_module(
        doc,
        "Karyawan",
        "Modul Karyawan digunakan untuk mengelola akun pengguna toko, termasuk nama, email, peran, status, tanggal bergabung, dan kata sandi.",
        "Owner dapat mengelola karyawan. Super Admin memiliki cakupan platform, sedangkan Cashier tidak mengakses modul ini.",
        [
            "Buka menu Karyawan.",
            "Klik Tambah Karyawan.",
            "Isi Nama Karyawan, Alamat Email, dan Kata Sandi.",
            "Pilih Peran: Kasir atau Pemilik (Owner).",
            "Pilih Status: Aktif atau Nonaktif.",
            "Klik Tambah Karyawan atau Simpan Perubahan.",
            "Pada mode edit, kosongkan field kata sandi jika tidak ingin mengubah kata sandi pengguna.",
        ],
        [
            "Nonaktifkan akun karyawan yang sudah tidak bertugas agar tidak dapat digunakan untuk login.",
            "Gunakan peran Cashier untuk staf kasir harian dan Owner hanya untuk pengguna yang boleh melihat laporan serta pengaturan.",
        ],
        16,
        "15-employees.png",
        "Daftar karyawan, peran, status, dan aksi",
    )

    doc.add_heading("Pengaturan", level=1)
    doc.add_paragraph("Modul Pengaturan berisi Profil Toko, Metode Pembayaran, dan Informasi Akun. Pengaturan ini memengaruhi identitas toko, pajak pada transaksi POS, logo struk, serta metode pembayaran yang dapat dipilih kasir.")
    add_figure(doc, 17, "16-settings.png", "Profil toko, metode pembayaran, dan informasi akun")
    add_steps(doc, [
        "Buka menu Pengaturan.",
        "Pada Profil Toko, lengkapi nama toko, email, nomor telepon, jam operasional, alamat, mata uang, tarif pajak, dan template struk.",
        "Klik Unggah logo untuk mengganti logo toko.",
        "Klik Simpan Perubahan untuk menyimpan profil toko.",
        "Pada Metode Pembayaran, aktifkan atau nonaktifkan metode pembayaran menggunakan switch.",
        "Periksa Informasi Akun untuk melihat paket layanan, status, ID tenant, dan tanggal pendaftaran.",
    ])
    add_note(doc, "Tips.", "Pastikan tarif pajak diperiksa sebelum POS digunakan. Perubahan tarif pajak akan memengaruhi perhitungan total transaksi berikutnya.", "success")

    add_module(
        doc,
        "Platform Admin",
        "Platform Admin digunakan oleh Super Admin untuk memantau tenant/toko, statistik platform, paket layanan, status tenant, jumlah pengguna, jumlah produk, total transaksi, dan GMV.",
        "Hanya Super Admin yang dapat mengakses halaman ini.",
        [
            "Login menggunakan akun Super Admin.",
            "Buka menu Platform Admin.",
            "Periksa statistik Total Toko, Total Pengguna, Total Transaksi, dan GMV Platform.",
            "Lihat kartu paket layanan untuk memahami batas produk, karyawan, dan fitur.",
            "Pada tabel Daftar Toko UMKM, ubah paket tenant melalui dropdown Paket bila diperlukan.",
            "Klik Tangguhkan untuk menonaktifkan tenant aktif, atau Aktifkan untuk mengaktifkan kembali tenant yang ditangguhkan.",
        ],
        [
            "Perubahan paket dan status tenant berdampak pada akses operasional toko.",
            "Gunakan tindakan tangguhkan hanya berdasarkan prosedur bisnis yang disepakati perusahaan.",
        ],
        18,
        "17-platform-admin.png",
        "Halaman Platform Admin untuk pengelolaan tenant",
    )

    doc.add_heading("Praktik Penggunaan yang Disarankan", level=1)
    doc.add_heading("Rutinitas Harian Kasir", level=2)
    add_bullets(doc, [
        "Login menggunakan akun pribadi.",
        "Periksa Dashboard untuk melihat kondisi awal transaksi dan stok.",
        "Gunakan POS untuk setiap penjualan dan pastikan pembayaran cukup sebelum checkout.",
        "Cetak atau cetak ulang struk bila pelanggan membutuhkan bukti transaksi.",
        "Logout setelah shift selesai.",
    ])
    doc.add_heading("Rutinitas Harian Owner", level=2)
    add_bullets(doc, [
        "Pantau stok menipis dari Dashboard dan Inventaris.",
        "Perbarui produk, harga, dan stok sebelum toko beroperasi.",
        "Cek Transaksi dan Laporan untuk memastikan penjualan tercatat wajar.",
        "Catat biaya operasional secara rutin agar laporan laba rugi lebih akurat.",
        "Kelola status karyawan sesuai perubahan staf.",
    ])
    doc.add_heading("Kontrol Data dan Keamanan", level=2)
    add_bullets(doc, [
        "Gunakan akun terpisah untuk setiap pengguna.",
        "Batasi peran Owner hanya kepada pengguna yang memang berwenang.",
        "Periksa ulang sebelum menghapus produk, kategori, pelanggan, pemasok, karyawan, atau biaya operasional.",
        "Lakukan refund hanya bila transaksi valid untuk dibatalkan.",
        "Jaga konsistensi data master agar laporan dan pencarian tetap akurat.",
    ])

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = closing.add_run("Akhir Dokumen")
    r.bold = True
    r.font.color.rgb = RGBColor(15, 118, 110)

    DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
